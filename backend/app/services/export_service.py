import io
from typing import List
from sqlalchemy.orm import Session

from app.models.question import Question
from app.models.answer import Answer
from app.models.project import Project


def export_to_docx(db: Session, project_id: str, include_unanswered: bool = False) -> bytes:
    """Generates a Word document with Q&A pairs. Returns bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise ValueError("Project not found")

    # Get questions with answers
    questions_query = db.query(Question).filter(
        Question.project_id == project_id,
        Question.status.in_(["answered", "in_progress", "approved", "skipped"] if include_unanswered else ["answered"]),
    ).order_by(Question.number)

    questions = questions_query.all()

    doc = Document()

    # Title
    title = doc.add_heading(project.name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if project.deadline:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Deadline: {project.deadline.strftime('%B %d, %Y')}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    current_section = None

    for question in questions:
        # Section heading
        if question.section and question.section != current_section:
            current_section = question.section
            doc.add_heading(question.section, level=1)

        # Question
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{question.number}. {question.text}")
        q_run.bold = True
        q_run.font.size = Pt(11)

        # Answer
        answer = db.query(Answer).filter(
            Answer.question_id == question.id,
            Answer.is_current == True,
        ).first()

        if answer:
            a_para = doc.add_paragraph()
            a_para.add_run(answer.content)
            a_para.style.font.size = Pt(10)
        else:
            a_para = doc.add_paragraph()
            run = a_para.add_run("[No answer yet]")
            run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            run.font.italic = True

        doc.add_paragraph()  # spacing

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
