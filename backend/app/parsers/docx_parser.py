from docx import Document
from typing import List, Dict, Any


def parse_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Parses DOCX, groups paragraphs by heading sections.
    Returns list of {page, section, text} dicts.
    """
    doc = Document(file_path)
    sections = []
    current_section = None
    current_texts = []
    fake_page = 1
    chars_per_page = 3000

    total_chars = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        is_heading = "Heading" in style_name

        if is_heading:
            if current_texts:
                sections.append({
                    "page": fake_page,
                    "section": current_section,
                    "text": "\n".join(current_texts),
                })
                total_chars += sum(len(t) for t in current_texts)
                fake_page = max(1, total_chars // chars_per_page + 1)
            current_section = text
            current_texts = []
        else:
            current_texts.append(text)

    if current_texts:
        sections.append({
            "page": fake_page,
            "section": current_section,
            "text": "\n".join(current_texts),
        })

    if not sections:
        # Fallback: return all text as one block
        all_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        if all_text:
            sections.append({"page": 1, "section": None, "text": all_text})

    return sections
